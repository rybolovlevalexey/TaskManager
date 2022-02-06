import sys
import datetime
import csv
import docx
from random import randint
from PyQt5.QtWidgets import QCheckBox
from PyQt5.QtWidgets import QApplication
from PyQt5.QtWidgets import QPushButton
from PyQt5.QtWidgets import QWidget
from PyQt5.QtWidgets import QLabel
from PyQt5.QtWidgets import QVBoxLayout
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtWidgets import QAction
from PyQt5.QtWidgets import QLineEdit
from PyQt5.QtWidgets import QPlainTextEdit
from PyQt5.QtWidgets import QDateTimeEdit
from PyQt5.QtWidgets import QComboBox
from PyQt5.QtWidgets import QTableWidget
from PyQt5.QtWidgets import QTableWidgetItem
from PyQt5.QtGui import QFont
from PyQt5.QtGui import QIcon
from PyQt5.QtGui import QColor


class InformationWidget(QWidget):
    def __init__(self):
        super(InformationWidget, self).__init__()
        self.setGeometry(700, 400, 500, 500)
        self.setWindowTitle('Информация о задаче')
        self.font = QFont()
        self.font.setPointSize(14)

        self.info_table = QTableWidget(self)
        self.info_table.move(25, 25)
        self.info_table.resize(450, 450)
        self.info_table.setColumnCount(2)
        self.info_table.setHorizontalHeaderItem(0, QTableWidgetItem(''))
        self.info_table.setHorizontalHeaderItem(1, QTableWidgetItem(''))
        self.info_table.setFont(self.font)
        self.info_table.setRowCount(5)
        self.info_table.setColumnWidth(0, 145)
        self.info_table.setColumnWidth(1, 280)
        self.info_table.setRowHeight(0, 80)
        self.info_table.setRowHeight(1, 80)
        self.info_table.setRowHeight(2, 80)
        self.info_table.setRowHeight(3, 80)
        self.info_table.setRowHeight(4, 80)

        self.info_table.setItem(0, 0, QTableWidgetItem('Название задачи'))
        self.info_table.setItem(1, 0, QTableWidgetItem('Подробнее'))
        self.info_table.setItem(2, 0, QTableWidgetItem('Дата и время начала выполнения задачи'))
        self.info_table.setItem(3, 0, QTableWidgetItem('Приоритет задачи'))
        self.info_table.setItem(4, 0, QTableWidgetItem('Результат'))

    def set_info_in_lbl(self, id):
        file = open('tasks.csv', encoding='UTF-8')
        reader = list(csv.reader(file, delimiter=';'))
        self.n = 0
        for self.elem in reader:
            if str(self.elem[0]) == str(id):
                self.info_table.setItem(0, 1, QTableWidgetItem(self.elem[1]))
                self.info_table.setItem(1, 1, QTableWidgetItem(self.elem[2]))
                self.info_table.setItem(2, 1,
                                        QTableWidgetItem(self.elem[3]))
                prioritet = self.elem[4]
                if prioritet == '1':
                    prioritet = 'Приоритет в порядке очереди'
                elif prioritet == '2':
                    prioritet = 'Средний приоритет'
                elif prioritet == '3':
                    prioritet = 'Высокий приоритет'
                else:
                    prioritet = 'Наивысший приоритет'

                self.info_table.setItem(3, 1, QTableWidgetItem(prioritet))
                self.info_table.setItem(4, 1, QTableWidgetItem(self.elem[-1]))

                if self.elem[-1] == 'Resolved':
                    self.setFixedSize(500, 650)
                    self.btn_not_solved = QPushButton(self)
                    self.btn_not_solved.setFont(self.font)
                    self.btn_not_solved.setText('Отправить задачу на доработку')
                    self.btn_not_solved.move(100, 550)
                    self.btn_not_solved.resize(300, 50)
                    self.btn_not_solved.clicked.connect(self.not_solved_pushed)

                break
            self.n += 1

    def not_solved_pushed(self):
        file = open('tasks.csv', encoding='UTF-8', newline='')
        reader = list(csv.reader(file, delimiter=';'))
        file.close()
        new_file = list()
        for i in range(len(reader)):
            elem = reader[i]
            if i == self.n:
                elem[-1] = 'NotResolved'
                new_file.append(elem)
            else:
                new_file.append(elem)
        file = open('tasks.csv', encoding='UTF-8', newline='', mode='w')
        writer = csv.writer(file, delimiter=';')
        writer.writerows(new_file)
        file.close()


class AboutWindow(QWidget):
    def __init__(self):
        super(AboutWindow, self).__init__()
        self.setGeometry(700, 400, 200, 200)
        self.setWindowTitle('О приложении')
        self.font1 = QFont()
        self.font1.setPointSize(14)
        self.font1.setBold(True)

        self.setLayout(QVBoxLayout(self))
        self.info = QLabel(self)
        self.info.setText('Данное приложение позволяет\n'
                          'систематизировать\n'
                          'все свои задачи и\n'
                          'просматривать информацию о них')
        self.info.setFont(self.font1)
        self.layout().addWidget(self.info)


class HowItWorkWindow(QWidget):
    def __init__(self):
        super(HowItWorkWindow, self).__init__()
        self.setGeometry(700, 400, 200, 200)
        self.setWindowTitle('Обучение')
        self.font1 = QFont()
        self.font1.setPointSize(14)
        self.font1.setBold(True)

        self.setLayout(QVBoxLayout(self))
        self.info = QLabel(self)
        self.info.setText('Нажмите кнопку "Добавить задачу"\n'
                          'для добавления новой задачи.\n'
                          'Нажмите кнопку "Просмотреть задачи"\n'
                          'для просмотра всех задач')
        self.info.setFont(self.font1)
        self.layout().addWidget(self.info)


class FeedbackWindow(QWidget):
    def __init__(self):
        super(FeedbackWindow, self).__init__()
        self.setGeometry(700, 400, 400, 250)
        self.setWindowTitle('Отзыв')
        self.lbl = QLabel(self)

        self.font1 = QFont()
        self.font1.setPointSize(14)
        self.font1.setBold(True)

        self.font2 = QFont()
        self.font2.setPointSize(14)

        self.font3 = QFont()
        self.font3.setPointSize(20)
        self.font3.setBold(True)

        self.info = QLabel(self)
        self.info.setText('Введите свой отзыв')
        self.info.setFont(self.font1)
        self.info.move(100, 25)

        self.input_feedback = QLineEdit(self)
        self.input_feedback.move(50, 75)
        self.input_feedback.resize(300, 50)
        self.input_feedback.setFont(self.font2)

        self.feedback_btn = QPushButton(self)
        self.feedback_btn.setText('Отправить отзыв')
        self.feedback_btn.setFont(self.font2)
        self.feedback_btn.resize(200, 40)
        self.feedback_btn.move(100, 175)
        self.feedback_btn.clicked.connect(self.btn_pushed)

    def btn_pushed(self):
        self.info.deleteLater()
        self.input_feedback.deleteLater()
        self.feedback_btn.deleteLater()

        self.lbl.setText('Спасибо за\n'
                         'отправленный\n'
                         'вами отзыв')
        self.lbl.setFont(self.font3)
        self.lbl.resize(250, 200)
        self.lbl.move(100, 25)


class DonatesWindow(QWidget):
    def __init__(self):
        super(DonatesWindow, self).__init__()
        self.setGeometry(700, 400, 200, 200)
        self.setWindowTitle('Донаты')
        self.font1 = QFont()
        self.font1.setPointSize(14)
        self.font1.setBold(True)

        self.setLayout(QVBoxLayout(self))
        self.info = QLabel(self)
        self.info.setText('Пожалуйста, поддержите автора:\n'
                          '5102 2945 2515 0848')
        self.info.setFont(self.font1)
        self.layout().addWidget(self.info)


class AuthorWindow(QWidget):
    def __init__(self):
        super(AuthorWindow, self).__init__()
        self.setGeometry(700, 400, 300, 150)
        self.setWindowTitle('Автор')
        self.font1 = QFont()
        self.font1.setPointSize(14)
        self.font1.setBold(True)

        self.info = QLabel(self)
        self.info.setText('@rybolovlevalexey')
        self.info.setFont(self.font1)
        self.info.move(60, 75)


class WidgetWatchingTasks(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Просмотр задач')
        self.setFixedSize(800, 600)
        self.move(500, 300)
        self.last_id = -1
        self.widg = InformationWidget()

        self.font = QFont()
        self.font.setPointSize(16)

        self.font_big = QFont()
        self.font_big.setPointSize(20)
        self.font_big.setBold(True)

        self.table = QTableWidget(self)
        self.table.setFont(self.font)
        self.table.move(50, 100)
        self.table.resize(700, 350)
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderItem(0, QTableWidgetItem('Название задачи'))
        self.table.setHorizontalHeaderItem(1, QTableWidgetItem('Подробнее'))
        self.table.setHorizontalHeaderItem(2, QTableWidgetItem('Начало выполнения'))
        self.table.setColumnWidth(0, 200)
        self.table.setColumnWidth(1, 245)
        self.table.setColumnWidth(2, 230)
        self.table.setColumnWidth(3, 0)
        self.table.setAlternatingRowColors(True)
        self.table.clicked.connect(self.table_pushed)

        self.btn_word = QPushButton(self)
        self.btn_word.resize(175, 75)
        self.btn_word.move(50, 500)
        self.btn_word.setFont(self.font)
        self.btn_word.setText('Конвертировать\n'
                              'в Word(docx)')
        self.btn_word.clicked.connect(self.word_pushed)

        self.btn_solved = QPushButton(self)
        self.btn_solved.resize(200, 75)
        self.btn_solved.move(300, 500)
        self.btn_solved.setFont(self.font_big)
        self.btn_solved.setText('Выполнено')
        self.btn_solved.clicked.connect(self.solved_pushed)

        self.btn_info = QPushButton(self)
        self.btn_info.move(575, 500)
        self.btn_info.resize(175, 75)
        self.btn_info.setFont(self.font)
        self.btn_info.setText('Просмотреть\n'
                              'информацию')
        self.btn_info.clicked.connect(self.info_pushed)

        self.btn_sort = QPushButton(self)
        self.btn_sort.move(150, 25)
        self.btn_sort.resize(200, 50)
        self.btn_sort.setFont(self.font)
        self.btn_sort.setText('Обновить')
        self.btn_sort.clicked.connect(self.sort_pushed)

        self.box_keys = QComboBox(self)
        self.box_keys.setFont(self.font)
        self.box_keys.move(400, 25)
        self.box_keys.resize(250, 50)
        self.box_keys.addItem('По дате добавления')
        self.box_keys.addItem('Только выполненные')
        self.box_keys.addItem('Только невыполненные')
        self.box_keys.addItem('По дате старта')
        self.box_keys.addItem('В алфавитном порядке по названию')
        self.box_keys.addItem('По приоритету')

        file = open('tasks.csv', encoding='UTF-8')
        reader = list(csv.reader(file, delimiter=';'))
        self.reader = list(reader)
        self.table.setRowCount(len(self.reader))
        for n in range(len(self.reader)):
            elem = self.reader[n]
            self.table.setItem(n, 0, QTableWidgetItem(elem[1]))
            self.table.setItem(n, 1, QTableWidgetItem(elem[2]))
            self.table.setItem(n, 2, QTableWidgetItem(elem[3]))
            self.table.setItem(n, 3, QTableWidgetItem(elem[0]))
            if elem[5] == 'NotResolved':
                self.table.item(n, 0).setBackground(QColor('red'))
            else:
                self.table.item(n, 0).setBackground(QColor('green'))

    def info_pushed(self):
        if self.last_id == -1:
            return
        self.widg = InformationWidget()
        self.widg.set_info_in_lbl(self.last_id)
        self.widg.show()

    def word_pushed(self):
        if self.last_id == -1:
            return
        document = docx.Document()
        document.add_heading('Отчёт о задаче(Issue report)', 0)
        paragraf = document.add_paragraph()

        file = open('tasks.csv', encoding='UTF-8')
        reader = list(csv.reader(file, delimiter=';'))
        for elem in reader:
            if elem[0] == self.last_id:
                title = elem[1]
                more = elem[2]
                dt_start = elem[3]
                prioritet = elem[-2]
                result = elem[-1]

                if prioritet == '1':
                    prioritet = 'Приоритет в порядке очереди'
                elif prioritet == '2':
                    prioritet = 'Средний приоритет'
                elif prioritet == '3':
                    prioritet = 'Высокий приоритет'
                else:
                    prioritet = 'Наивысший приоритет'

                if result == 'Resolved':
                    result = 'Выполнено успешно'
                else:
                    result = 'Ещё не выполнено'

                s = paragraf.add_run(f'Название задачи(title):\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(20)

                s = paragraf.add_run(f'{title}\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(16)

                s = paragraf.add_run(f'Подробнее(more detailed):\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(20)

                s = paragraf.add_run(f'{more}\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(16)

                s = paragraf.add_run(
                    f'Дата и время начала выполнения задачи(date and time when the task started):\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(20)

                s = paragraf.add_run(f'{dt_start}\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(16)

                s = paragraf.add_run(f'Приоритет задачи(priority of the task):\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(20)

                s = paragraf.add_run(f'{prioritet}\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(16)

                s = paragraf.add_run(f'Результат(result):\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(20)

                s = paragraf.add_run(f'{result}\n')
                s.font.name = 'Arial'
                s.font.size = docx.shared.Pt(16)

                document.save('Issue report.docx')
                break

    def table_pushed(self):
        for elem in self.table.selectedItems():
            self.last_id = str(self.table.item(elem.row(), 3).text())

    def solved_pushed(self):
        if self.last_id == -1:
            return
        file = open('tasks.csv', encoding='UTF-8')
        reader = list(csv.reader(file, delimiter=';'))
        file1 = open('tasks.csv', encoding='UTF-8', mode='w', newline='')
        writer = csv.writer(file1, delimiter=';', quoting=csv.QUOTE_MINIMAL)
        self.reader = list(reader)
        n = 0
        new_reader = list()
        for n in range(len(self.reader)):
            if str(self.reader[n][0]) == self.last_id and self.reader[n][-1] != 'Resolved':
                line = self.reader[n]
                line[-1] = 'Resolved'
                new_reader.append(list(line))
                break
            elif str(self.reader[n][0]) != self.last_id:
                new_reader.append(self.reader[n])
        writer.writerows(new_reader)
        file.close()
        file1.close()

    def sort_pushed(self):
        condition = self.box_keys.currentText()
        if condition == 'По дате добавления':
            self.table.setRowCount(0)
            file = open('tasks.csv', encoding='UTF-8')
            reader = list(csv.reader(file, delimiter=';'))
            self.reader = list(reader)
            self.table.setRowCount(len(self.reader))
            for n in range(len(self.reader)):
                elem = self.reader[n]
                self.table.setItem(n, 0, QTableWidgetItem(elem[1]))
                self.table.setItem(n, 1, QTableWidgetItem(elem[2]))
                self.table.setItem(n, 2, QTableWidgetItem(elem[3]))
                self.table.setItem(n, 3, QTableWidgetItem(elem[0]))
                if elem[5] == 'NotResolved':
                    self.table.item(n, 0).setBackground(QColor('red'))
                else:
                    self.table.item(n, 0).setBackground(QColor('green'))
            file.close()
        elif condition == 'Только выполненные':
            self.table.setRowCount(0)
            file = open('tasks.csv', encoding='UTF-8')
            reader = list(csv.reader(file, delimiter=';'))
            self.reader = list(reader)
            colvo_row = 1
            for n in range(len(self.reader)):
                elem = self.reader[n]
                if elem[5] == 'Resolved':
                    self.table.setRowCount(colvo_row)
                    self.table.setItem(colvo_row - 1, 0, QTableWidgetItem(elem[1]))
                    self.table.setItem(colvo_row - 1, 1, QTableWidgetItem(elem[2]))
                    self.table.setItem(colvo_row - 1, 2, QTableWidgetItem(elem[3]))
                    self.table.setItem(colvo_row - 1, 3, QTableWidgetItem(elem[0]))
                    colvo_row += 1
            file.close()
        elif condition == 'Только невыполненные':
            self.table.setRowCount(0)
            file = open('tasks.csv', encoding='UTF-8')
            reader = list(csv.reader(file, delimiter=';'))
            self.reader = list(reader)
            row_colvo = 1
            for n in range(len(self.reader)):
                elem = self.reader[n]
                if elem[5] == 'NotResolved':
                    self.table.setRowCount(row_colvo)
                    self.table.setItem(row_colvo - 1, 0, QTableWidgetItem(elem[1]))
                    self.table.setItem(row_colvo - 1, 1, QTableWidgetItem(elem[2]))
                    self.table.setItem(row_colvo - 1, 2, QTableWidgetItem(elem[3]))
                    self.table.setItem(row_colvo - 1, 3, QTableWidgetItem(elem[0]))
                    row_colvo += 1
            file.close()
        elif condition == 'По дате старта':
            self.table.setRowCount(0)
            file = open('tasks.csv', encoding='UTF-8')
            reader = list(csv.reader(file, delimiter=';'))
            self.reader = sorted(list(reader), key=lambda x: x[3])
            for n in range(len(self.reader)):
                elem = self.reader[n]
                self.table.setRowCount(n + 1)
                self.table.setItem(n, 0, QTableWidgetItem(elem[1]))
                self.table.setItem(n, 1, QTableWidgetItem(elem[2]))
                self.table.setItem(n, 2, QTableWidgetItem(elem[3]))
                self.table.setItem(n, 3, QTableWidgetItem(elem[0]))
                if elem[5] == 'NotResolved':
                    self.table.item(n, 0).setBackground(QColor('red'))
                else:
                    self.table.item(n, 0).setBackground(QColor('green'))
            file.close()
        elif condition == 'В алфавитном порядке по названию':
            self.table.setRowCount(0)
            file = open('tasks.csv', encoding='UTF-8')
            reader = list(csv.reader(file, delimiter=';'))
            self.reader = sorted(list(reader), key=lambda x: x[1])
            for n in range(len(self.reader)):
                elem = self.reader[n]
                self.table.setRowCount(n + 1)
                self.table.setItem(n, 0, QTableWidgetItem(elem[1]))
                self.table.setItem(n, 1, QTableWidgetItem(elem[2]))
                self.table.setItem(n, 2, QTableWidgetItem(elem[3]))
                self.table.setItem(n, 3, QTableWidgetItem(elem[0]))
                if elem[5] == 'NotResolved':
                    self.table.item(n, 0).setBackground(QColor('red'))
                else:
                    self.table.item(n, 0).setBackground(QColor('green'))
        elif condition == 'По приоритету':
            self.table.setRowCount(0)
            file = open('tasks.csv', encoding='UTF-8')
            reader = list(csv.reader(file, delimiter=';'))
            self.reader = sorted(list(reader), key=lambda x: x[-2])[::-1]
            self.table.setRowCount(len(self.reader))
            for n in range(len(self.reader)):
                elem = self.reader[n]
                self.table.setItem(n, 0, QTableWidgetItem(elem[1]))
                self.table.setItem(n, 1, QTableWidgetItem(elem[2]))
                self.table.setItem(n, 2, QTableWidgetItem(elem[3]))
                self.table.setItem(n, 3, QTableWidgetItem(elem[0]))
                if elem[5] == 'NotResolved':
                    self.table.item(n, 0).setBackground(QColor('red'))
                else:
                    self.table.item(n, 0).setBackground(QColor('green'))
            file.close()


class WidgetAppendNewTask(QWidget):
    def __init__(self):
        super().__init__()
        self.setFixedSize(800, 600)
        self.move(500, 300)
        self.setWindowTitle('Добавление новой задачи')

        self.font1 = QFont()
        self.font1.setPointSize(16)

        self.lbl1 = QLabel(self)
        self.lbl1.move(50, 50)
        self.lbl1.resize(350, 50)
        self.lbl1.setFont(self.font1)
        self.lbl1.setText('Название задачи')

        self.line_name = QLineEdit(self)
        self.line_name.move(400, 50)
        self.line_name.resize(350, 50)
        self.line_name.setFont(self.font1)

        self.lbl2 = QLabel(self)
        self.lbl2.move(50, 125)
        self.lbl2.resize(350, 50)
        self.lbl2.setFont(self.font1)
        self.lbl2.setText('Подробнее')

        self.text_of_task = QPlainTextEdit(self)
        self.text_of_task.move(400, 125)
        self.text_of_task.resize(350, 150)
        self.text_of_task.setFont(self.font1)

        self.lbl3 = QLabel(self)
        self.lbl3.move(50, 300)
        self.lbl3.resize(350, 50)
        self.lbl3.setText('Укажите дату и время\n'
                          'начала выполнения задачи')
        self.lbl3.setFont(self.font1)

        self.start_check = QCheckBox(self)
        self.start_check.move(400, 300)
        self.start_check.resize(20, 50)
        self.start_check.setFont(self.font1)
        self.start_check.clicked.connect(self.start_check_triggered)

        self.dt_start = QDateTimeEdit(self)
        self.dt_start.move(430, 300)
        self.dt_start.resize(320, 50)
        self.dt_start.setFont(self.font1)
        self.dt_start.setDateTime(datetime.datetime.now())
        self.dt_start.setDisabled(True)

        self.lbl4 = QLabel(self)
        self.lbl4.move(50, 400)
        self.lbl4.resize(350, 50)
        self.lbl4.setText('Установите приоритет задачи')
        self.lbl4.setFont(self.font1)

        self.prioritet = QComboBox(self)
        self.prioritet.move(400, 400)
        self.prioritet.resize(350, 50)
        self.prioritet.setFont(self.font1)
        self.prioritet.addItem('Приоритет в порядке очереди')
        self.prioritet.addItem('Средний приоритет')
        self.prioritet.addItem('Высокий приоритет')
        self.prioritet.addItem('Наивысший приоритет')

        self.btn_back = QPushButton(self)
        self.btn_back.move(0, 0)
        self.btn_back.resize(50, 25)
        self.btn_back.setFont(self.font1)
        self.btn_back.setText('<-')
        self.btn_back.clicked.connect(self.back_pushed)

        self.btn_ready = QPushButton(self)
        self.btn_ready.move(250, 500)
        self.btn_ready.resize(300, 50)
        self.btn_ready.setFont(self.font1)
        self.btn_ready.setText('Сохранить задачу')
        self.btn_ready.clicked.connect(self.ready_pushed)

        self.info_lbl = QLabel(self)
        self.info_lbl.move(150, 0)
        self.info_lbl.setFont(self.font1)
        self.info_lbl.setText('')
        self.info_lbl.setStyleSheet('color: red')

    def back_pushed(self):
        if self.line_name.text() == '':
            self.close()
        else:
            self.info_lbl.setText('Очистите поле "Название задачи" или сохраните задачу')
            self.info_lbl.resize(555, 25)

    def mousePressEvent(self, event):
        if event.button() == 8 and self.line_name.text() == '':
            self.close()
        else:
            self.info_lbl.setText('Очистите поле "Название задачи" или сохраните задачу')
            self.info_lbl.resize(555, 25)

    def start_check_triggered(self):
        if self.start_check.isChecked():
            self.dt_start.setDisabled(False)
        else:
            self.dt_start.setDisabled(True)

    def ready_pushed(self):
        file = open('tasks.csv', encoding='UTF-8', mode='a', newline='')
        writer = csv.writer(file, delimiter=';',
                            quoting=csv.QUOTE_MINIMAL)
        name = self.line_name.text()
        more = self.text_of_task.toPlainText()
        prioritet = self.prioritet.currentText()
        if self.start_check.isChecked():
            start = self.dt_start.dateTime()
        else:
            start = 'None'
        dt = str(start.toPyDateTime())
        dt = dt[:dt.index('.')]
        dt = dt[:-3]

        if prioritet == 'Приоритет в порядке очереди':
            prioritet = 1
        elif prioritet == 'Средний приоритет':
            prioritet = 2
        elif prioritet == 'Высокий приоритет':
            prioritet = 3
        else:
            prioritet = 4

        row = [str(randint(10000, 99999)), name, more, dt, str(prioritet), 'NotResolved']
        writer.writerow(row)
        self.info_lbl.setText('Задача успешно сохранена')
        self.info_lbl.resize(555, 25)
        self.close()


class Window(QMainWindow):
    def __init__(self):
        super(Window, self).__init__()
        self.setGeometry(500, 300, 800, 600)
        self.setFixedSize(800, 600)
        self.setWindowIcon(QIcon('иконка8.png'))
        self.setWindowTitle('Task Manager')

        self.font1 = QFont()
        self.font1.setPointSize(20)

        self.btn_new_task = QPushButton(self)
        self.btn_new_task.setText('Добавить задачу')
        self.btn_new_task.resize(300, 100)
        self.btn_new_task.move(250, 150)
        self.btn_new_task.setFont(self.font1)
        self.btn_new_task.clicked.connect(self.new_task_pushed)

        self.btn_watch_tasks = QPushButton(self)
        self.btn_watch_tasks.setText('Просмотреть задачи')
        self.btn_watch_tasks.resize(300, 100)
        self.btn_watch_tasks.move(250, 350)
        self.btn_watch_tasks.setFont(self.font1)
        self.btn_watch_tasks.clicked.connect(self.watch_tasks_pushed)

        self.about_action = QAction(self)
        self.about_action.triggered.connect(self.about_show)
        self.about_action.setText('О приложении')
        self.menuBar().addAction(self.about_action)

        self.how_it_work = QAction(self)
        self.how_it_work.triggered.connect(self.show_how_it_work)
        self.how_it_work.setText('Обучение')
        self.menuBar().addAction(self.how_it_work)

        self.feedback = QAction(self)
        self.feedback.triggered.connect(self.feedback_show)
        self.feedback.setText('Отзыв')
        self.menuBar().addAction(self.feedback)

        self.donates = QAction(self)
        self.donates.triggered.connect(self.donates_show)
        self.donates.setText('Донаты')
        self.menuBar().addAction(self.donates)

        self.author = QAction(self)
        self.author.triggered.connect(self.author_show)
        self.author.setText('Автор')
        self.menuBar().addAction(self.author)

        self.about_window = AboutWindow()
        self.how_it_work_window = HowItWorkWindow()
        self.donates_window = DonatesWindow()
        self.author_window = AuthorWindow()
        self.feedback_window = FeedbackWindow()

        self.new_task_window = WidgetAppendNewTask()
        self.watch_tasks_window = WidgetWatchingTasks()

    def about_show(self):
        self.about_window.show()

    def show_how_it_work(self):
        self.how_it_work_window.show()

    def feedback_show(self):
        self.feedback_window = FeedbackWindow()
        self.feedback_window.show()

    def donates_show(self):
        self.donates_window.show()

    def author_show(self):
        self.author_window.show()

    def new_task_pushed(self):
        self.new_task_window = WidgetAppendNewTask()
        self.new_task_window.show()

    def watch_tasks_pushed(self):
        self.watch_tasks_window = WidgetWatchingTasks()
        self.watch_tasks_window.show()


def except_hook(cls, exception, traceback):
    sys.__excepthook__(cls, exception, traceback)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Window()
    ex.show()
    sys.excepthook = except_hook
    sys.exit(app.exec())
